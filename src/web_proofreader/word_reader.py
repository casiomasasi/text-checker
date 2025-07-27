"""
Wordファイル読み込み機能

.docx, .doc ファイルからテキストを抽出し、
校正処理用の形式に変換する。
"""

import os
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Union
from docx import Document
import zipfile
import xml.etree.ElementTree as ET


class WordReader:
    """Wordファイル読み込み処理クラス"""
    
    SUPPORTED_EXTENSIONS = {'.docx', '.doc'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        self.current_file_path: Optional[Path] = None
        
    def validate_file(self, file_path: Union[str, Path]) -> Dict[str, Union[bool, str]]:
        """
        ファイルの妥当性をチェック
        
        Args:
            file_path: ファイルパス
            
        Returns:
            検証結果の辞書
        """
        file_path = Path(file_path)
        
        # ファイル存在チェック
        if not file_path.exists():
            return {
                'valid': False,
                'error': f'ファイルが見つかりません: {file_path}'
            }
        
        # 拡張子チェック
        if file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            return {
                'valid': False,
                'error': f'サポートされていないファイル形式: {file_path.suffix}'
            }
        
        # ファイルサイズチェック
        file_size = file_path.stat().st_size
        if file_size > self.MAX_FILE_SIZE:
            return {
                'valid': False,
                'error': f'ファイルサイズが大きすぎます: {file_size / 1024 / 1024:.1f}MB (上限: {self.MAX_FILE_SIZE / 1024 / 1024}MB)'
            }
        
        # .docxファイルの構造チェック
        if file_path.suffix.lower() == '.docx':
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    # 必要なファイルが存在するかチェック
                    required_files = ['word/document.xml', '[Content_Types].xml']
                    for required_file in required_files:
                        if required_file not in zip_file.namelist():
                            return {
                                'valid': False,
                                'error': '無効なWordファイル形式です'
                            }
            except zipfile.BadZipFile:
                return {
                    'valid': False,
                    'error': 'ファイルが破損しているか、有効なWordファイルではありません'
                }
        
        return {'valid': True}
    
    def read_docx(self, file_path: Union[str, Path]) -> Dict[str, Union[str, List[str], Dict]]:
        """
        .docxファイルを読み込み
        
        Args:
            file_path: .docxファイルのパス
            
        Returns:
            読み込み結果の辞書
        """
        file_path = Path(file_path)
        
        try:
            # python-docxを使用してテキストを抽出
            doc = Document(file_path)
            
            paragraphs = []
            full_text = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # 空でない段落のみ追加
                    paragraphs.append(text)
                    full_text.append(text)
            
            # メタデータの取得
            core_props = doc.core_properties
            
            metadata = {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'extension': file_path.suffix.lower(),
                'paragraph_count': len(paragraphs),
                'char_count': sum(len(p) for p in paragraphs),
                'word_count': sum(len(p.split()) for p in paragraphs),
                'title': core_props.title or '',
                'author': core_props.author or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else ''
            }
            
            return {
                'content': '\\n\\n'.join(full_text),
                'paragraphs': paragraphs,
                'metadata': metadata
            }
            
        except Exception as e:
            raise Exception(f"docxファイル読み込みエラー: {e}")
    
    def read_doc_fallback(self, file_path: Union[str, Path]) -> Dict[str, Union[str, List[str], Dict]]:
        """
        .docファイルの読み込み（フォールバック処理）
        
        Args:
            file_path: .docファイルのパス
            
        Returns:
            読み込み結果の辞書
        """
        file_path = Path(file_path)
        
        # .docファイルは複雑な形式のため、基本的な情報のみ提供
        metadata = {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_size': file_path.stat().st_size,
            'extension': file_path.suffix.lower(),
            'note': '.docファイルは.docx形式での保存を推奨します'
        }
        
        return {
            'content': '',
            'paragraphs': [],
            'metadata': metadata,
            'error': '.docファイルは現在サポートされていません。.docx形式で保存し直してからアップロードしてください。'
        }
    
    def read_file(self, file_path: Union[str, Path]) -> Dict[str, Union[str, List[str], Dict]]:
        """
        Wordファイルを読み込み
        
        Args:
            file_path: 読み込むファイルのパス
            
        Returns:
            読み込み結果の辞書
        """
        file_path = Path(file_path)
        self.current_file_path = file_path
        
        # ファイル検証
        validation_result = self.validate_file(file_path)
        if not validation_result['valid']:
            raise ValueError(validation_result['error'])
        
        # 拡張子に応じた処理
        if file_path.suffix.lower() == '.docx':
            return self.read_docx(file_path)
        elif file_path.suffix.lower() == '.doc':
            return self.read_doc_fallback(file_path)
        else:
            raise ValueError(f"サポートされていないファイル形式: {file_path.suffix}")
    
    def extract_text_with_positions(self, file_path: Union[str, Path]) -> Dict[str, Union[str, List[Dict]]]:
        """
        テキストを位置情報付きで抽出
        
        Args:
            file_path: ファイルパス
            
        Returns:
            位置情報付きテキストデータ
        """
        file_path = Path(file_path)
        
        if file_path.suffix.lower() != '.docx':
            raise ValueError("位置情報付き抽出は.docxファイルのみサポートしています")
        
        try:
            doc = Document(file_path)
            
            text_elements = []
            current_position = 0
            
            for paragraph_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text
                if text.strip():  # 空でない段落のみ処理
                    text_elements.append({
                        'type': 'paragraph',
                        'index': paragraph_idx,
                        'text': text,
                        'start_position': current_position,
                        'end_position': current_position + len(text),
                        'length': len(text)
                    })
                    current_position += len(text) + 2  # 段落間の改行を考慮
            
            full_text = '\\n\\n'.join([elem['text'] for elem in text_elements])
            
            return {
                'full_text': full_text,
                'elements': text_elements,
                'total_length': current_position
            }
            
        except Exception as e:
            raise Exception(f"位置情報付きテキスト抽出エラー: {e}")
    
    def get_file_info(self, file_path: Union[str, Path]) -> Dict[str, Union[str, int, bool]]:
        """
        ファイル情報を取得（読み込まずに）
        
        Args:
            file_path: ファイルパス
            
        Returns:
            ファイル情報の辞書
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"ファイルが見つかりません: {file_path}")
        
        stat = file_path.stat()
        validation_result = self.validate_file(file_path)
        
        return {
            'file_path': str(file_path),
            'file_name': file_path.name,
            'file_size': stat.st_size,
            'extension': file_path.suffix.lower(),
            'is_supported': file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS,
            'is_valid': validation_result['valid'],
            'validation_error': validation_result.get('error', ''),
            'size_mb': round(stat.st_size / 1024 / 1024, 2)
        }
    
    def create_temp_file(self, uploaded_file) -> str:
        """
        アップロードされたファイルを一時ファイルとして保存
        
        Args:
            uploaded_file: Flaskのファイルアップロードオブジェクト
            
        Returns:
            一時ファイルのパス
        """
        # ファイル名の安全性チェック
        filename = uploaded_file.filename
        if not filename:
            raise ValueError("ファイル名が指定されていません")
        
        # 安全なファイル名に変換
        safe_filename = self._make_safe_filename(filename)
        
        # 一時ファイル作成
        temp_dir = tempfile.mkdtemp()
        temp_file_path = os.path.join(temp_dir, safe_filename)
        
        # ファイル保存
        uploaded_file.save(temp_file_path)
        
        return temp_file_path
    
    def _make_safe_filename(self, filename: str) -> str:
        """
        安全なファイル名に変換
        
        Args:
            filename: 元のファイル名
            
        Returns:
            安全なファイル名
        """
        import re
        import unicodedata
        
        # Unicodeの正規化
        filename = unicodedata.normalize('NFKC', filename)
        
        # 危険な文字を除去
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        
        # 先頭・末尾の空白とピリオドを除去
        filename = filename.strip(' .')
        
        # 空のファイル名の場合はデフォルト名を使用
        if not filename:
            filename = 'uploaded_document.docx'
        
        return filename
    
    def cleanup_temp_file(self, temp_file_path: str):
        """
        一時ファイルを削除
        
        Args:
            temp_file_path: 一時ファイルのパス
        """
        try:
            file_path = Path(temp_file_path)
            if file_path.exists():
                file_path.unlink()
            
            # 一時ディレクトリも削除
            temp_dir = file_path.parent
            if temp_dir.exists() and temp_dir.name.startswith('tmp'):
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception as e:
            print(f"一時ファイル削除エラー: {e}")